import json
from io import BytesIO
from docx import Document
from docx.shared import Pt


def generate_prescription_doc(medicines: list[str]) -> BytesIO:
    # Load replacements
    with open("./Templates/data.json", "r", encoding="utf-8") as f:
        replacements = json.load(f)

    doc = Document("./Templates/TEMPLATE.docx")

    def replace_in_runs(runs, replacements):
        for run in runs:
            for tag, value in replacements.items():
                if tag in run.text:
                    run.text = run.text.replace(tag, value)

    # Replace placeholders
    for paragraph in doc.paragraphs:
        replace_in_runs(paragraph.runs, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_runs(paragraph.runs, replacements)

    # Add medicines section
    # Add medicines section
    doc.add_paragraph("")

    heading = doc.add_paragraph("Medicines Prescribed")
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(14)

    for med in medicines:
        p = doc.add_paragraph()
        run = p.add_run(f"• {med}")
        run.font.size = Pt(12)

    # Save to memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
